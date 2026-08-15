from __future__ import annotations

from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "SOXL_TQQQ Rotation with cash.docx"


def set_cell_text(cell, text: object, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(document: Document, frame: pd.DataFrame, widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    header = table.rows[0].cells
    for i, column in enumerate(frame.columns):
        set_cell_text(header[i], column, bold=True)
        header[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = header[i]._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4D78")
        shading.append(shd)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    document.add_paragraph()


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def pct(value: object) -> str:
    return f"{float(value):,.2f}%"


def main() -> None:
    cash_summary = pd.read_csv(ROOT / "SOXL_TQQQ_Cash_Regime_Summary.csv")
    comparison_2020 = pd.read_csv(ROOT / "SOXL_TQQQ_2020_To_Date_Summary.csv")
    annual_2020 = pd.read_csv(ROOT / "SOXL_TQQQ_2020_To_Date_Annual.csv")
    full_row = cash_summary[cash_summary["case"] == "Best balanced: DD <= 55%, 2022 loss <= 35%"].iloc[0]

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 77, 120)
    styles["Heading 2"].font.name = "Aptos"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.color.rgb = RGBColor(31, 77, 120)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SOXL/TQQQ Rotation with cash")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Daily scanner specification, historical performance, and automation notes")
    run.italic = True
    run.font.size = Pt(10)
    document.add_paragraph()

    document.add_heading("Executive Takeaway", level=1)
    document.add_paragraph(
        "The cash-balanced strategy keeps the high-upside SOXL/TQQQ rotation framework but adds a simple risk-off rule. "
        "It targets SOXL or TQQQ only when the selected ETF or QQQ confirms trend, and otherwise moves to cash. "
        "The result is lower drawdown and materially better bear-market behavior than the always-invested base rotation or the DCA return variant."
    )

    takeaway = pd.DataFrame(
        [
            ["Full test return", pct(full_row["net_return_pct"])],
            ["Full test max drawdown", pct(full_row["max_drawdown_pct"])],
            ["2022 return", pct(full_row["return_2022_pct"])],
            ["2022 drawdown", pct(full_row["drawdown_2022_pct"])],
            ["Cash days", pct(full_row["cash_days_pct"])],
            ["DCA", "No"],
        ],
        columns=["Metric", "Value"],
    )
    add_table(document, takeaway, [2.2, 2.4])

    document.add_heading("Trading Rules", level=1)
    document.add_heading("Base rotation", level=2)
    add_bullet(document, "Universe: SOXL and TQQQ only.")
    add_bullet(document, "Timeframe: daily bars, evaluated after the daily close.")
    add_bullet(document, "Relative momentum score: 63-day return minus 0.5 x annualized volatility, skipping the most recent 10 trading days.")
    add_bullet(document, "Rebalance cadence: monthly; a 5% score hysteresis is used to avoid small noisy switches.")
    add_bullet(document, "Trend preference: if the selected ETF is below its SMA50 and the alternate ETF is above its SMA50, prefer the alternate ETF.")

    document.add_heading("Cash filter", level=2)
    add_bullet(document, "Target can be SOXL, TQQQ, or CASH.")
    add_bullet(document, "Exit to cash when both the selected ETF and QQQ are below SMA150.")
    add_bullet(document, "Re-enter when either the selected ETF or QQQ is above SMA150 + 1%.")
    add_bullet(document, "Position size: 1x target position. The balanced cash version does not use DCA or leverage.")

    document.add_heading("Entry and Exit", level=1)
    entry_exit = pd.DataFrame(
        [
            ["Entry from cash", "Buy the current selected ETF when either the selected ETF or QQQ closes above SMA150 + 1%."],
            ["Rotation", "Switch between SOXL and TQQQ when the base monthly relative-momentum target changes and the cash filter remains risk-on."],
            ["Exit to cash", "Sell SOXL/TQQQ when both the selected ETF and QQQ are below SMA150."],
            ["Hold", "If the account already holds the target ETF, do nothing."],
            ["Execution window", "Automation runs at 4:10 PM ET and submits Alpaca DAY limit orders with extended_hours=True when the target is not aligned."],
        ],
        columns=["Action", "Rule"],
    )
    add_table(document, entry_exit, [1.6, 4.7])

    document.add_heading("Historical Performance", level=1)
    perf_full = pd.DataFrame(
        [
            ["New cash balanced strategy", pct(full_row["net_return_pct"]), pct(full_row["cagr_pct"]), pct(full_row["max_drawdown_pct"]), pct(full_row["return_2022_pct"])],
            ["Base SOXL/TQQQ rotation", pct(full_row["base_rotation_return_pct"]), "75.63%", pct(full_row["base_rotation_max_drawdown_pct"]), "-76.80%"],
            ["Verified DCA return strategy", "2,823,010.49%", "90.50%", "-65.54%", "-62.54%"],
            ["SOXL-only strategy", pct(full_row["soxl_only_return_pct"]), "46.02%", pct(full_row["soxl_only_max_drawdown_pct"]), "-57.81%"],
        ],
        columns=["Strategy", "Full return", "CAGR", "Max drawdown", "2022 return"],
    )
    add_table(document, perf_full, [2.2, 1.1, 0.9, 1.1, 1.1])

    document.add_heading("2020-to-date comparison", level=2)
    summary = comparison_2020.copy()
    summary["Cumulative return %"] = summary["Cumulative return %"].map(lambda value: f"{value:,.2f}%")
    summary["Max drawdown %"] = summary["Max drawdown %"].map(lambda value: f"{value:,.2f}%")
    summary["Ending growth of $1"] = summary["Ending growth of $1"].map(lambda value: f"{value:,.4f}")
    add_table(document, summary, [2.3, 1.6, 1.2, 1.2, 1.1])

    document.add_heading("Annual return and drawdown, 2020-to-date", level=2)
    annual = pd.DataFrame(
        {
            "Year": annual_2020["Year"],
            "Base return": annual_2020["Base SOXL/TQQQ rotation annual return %"].map(lambda v: f"{v:,.2f}%"),
            "Base DD": annual_2020["Base SOXL/TQQQ rotation max drawdown %"].map(lambda v: f"{v:,.2f}%"),
            "DCA return": annual_2020["Verified DCA strategy annual return %"].map(lambda v: f"{v:,.2f}%"),
            "DCA DD": annual_2020["Verified DCA strategy max drawdown %"].map(lambda v: f"{v:,.2f}%"),
            "Cash return": annual_2020["New cash balanced strategy annual return %"].map(lambda v: f"{v:,.2f}%"),
            "Cash DD": annual_2020["New cash balanced strategy max drawdown %"].map(lambda v: f"{v:,.2f}%"),
        }
    )
    add_table(document, annual, [0.55, 0.85, 0.75, 0.85, 0.75, 0.85, 0.75])

    document.add_heading("Why This Is Better Than Base and DCA Return", level=1)
    add_bullet(document, "Better bear-market behavior: in 2022 the cash strategy lost -27.73%, compared with -76.80% for the base rotation and -62.54% for the verified DCA return strategy.")
    add_bullet(document, "Lower structural drawdown: full-test max drawdown improved to -52.16%, versus -79.24% for base rotation and -65.54% for the DCA strategy.")
    add_bullet(document, "Cleaner risk logic: DCA adds exposure during declines; the cash strategy reduces exposure when both the selected ETF and QQQ fail trend.")
    add_bullet(document, "Still high-return: from 2020 to date, the cash strategy compounded $1 to $106.70, close to the DCA strategy's $113.96 while carrying much lower drawdown.")
    add_bullet(document, "Operationally simpler: no fractional exposure ladder or average-cost bookkeeping is required; the scanner targets SOXL, TQQQ, or CASH.")

    document.add_heading("Scanner and Automation", level=1)
    add_bullet(document, "Scanner script: scripts/soxl_tqqq_cash_signal_scanner.py")
    add_bullet(document, "Agent: agents/SOXL_TQQQ Rotation with cash")
    add_bullet(document, "The scanner loads Alpaca credentials from .env.alpaca via --env-file .env.alpaca.")
    add_bullet(document, "Automation ID: soxl-tqqq-rotation-with-cash-daily-scanner.")
    add_bullet(document, "Scheduled command: powershell -ExecutionPolicy Bypass -File scripts/run_soxl_tqqq_cash_daily_scanner_task.ps1 --agent \"SOXL/TQQQ Rotation with cash\" --env-file .env.alpaca --alpaca --execute --extended-hours --limit-offset-pct 0 --target-notional 10000 --qty 1 --email-to skydiver1118@gmail.com")
    add_bullet(document, "When target is CASH, the scanner sells SOXL/TQQQ positions. When target is SOXL or TQQQ, it sells the non-target ETF first and buys approximately $10,000 of the target ETF if not already held.")
    add_bullet(document, "The command defaults to Alpaca paper trading unless --live is added.")

    section = document.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.text = "Generated by Codex - local yfinance backtest; no commissions, slippage, taxes, or execution friction included."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        document.save(OUT)
        print(OUT)
    except PermissionError:
        unlocked_out = ROOT / f"SOXL_TQQQ Rotation with cash regenerated {datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        document.save(unlocked_out)
        print(unlocked_out)


if __name__ == "__main__":
    main()
