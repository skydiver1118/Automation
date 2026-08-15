from __future__ import annotations

from collections import Counter
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
RESULTS_XLSX = REPORT_DIR / "smh_historical_components_momentum_is2010_2019_oos2020_2026ytd.xlsx"
PRICE_CACHE = ROOT / "data/smh_components/historical_smh_open_close_2008-12-01_2026-05-24.csv"
OUT_DOCX = REPORT_DIR / "SMH_Historical_Momentum_Top2_L252_SMA100_DCA1_Report.docx"

STRATEGY = "SMH_HIST_PIT Top2 L252 S0 smh_sma100 DCA1"
TITLE = "SMH Historical Components Momentum Strategy"
SUBTITLE = "Top2 L252 S0 SMA100 DCA1 - Monthly Point-in-Time Backtest"
AS_OF = date(2026, 5, 24)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2EF"


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "DADCE0") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def paragraph_border_bottom(paragraph, color: str = "2E74B5", size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def set_table_indent(table, dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    ind = tbl_pr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), str(dxa))
    ind.set(qn("w:type"), "dxa")


def format_percent(value: float) -> str:
    return "" if pd.isna(value) else f"{value:.2%}"


def format_float(value: float) -> str:
    return "" if pd.isna(value) else f"{value:.2f}"


def metric_series(returns: pd.Series) -> dict[str, float]:
    clean = returns.dropna()
    equity = (1.0 + clean).cumprod()
    total_return = float(equity.iloc[-1] - 1.0)
    years = len(clean) / 12.0
    cagr = float(equity.iloc[-1] ** (1.0 / years) - 1.0)
    curve = pd.concat([pd.Series([1.0]), equity.reset_index(drop=True)], ignore_index=True)
    drawdown = curve / curve.cummax() - 1.0
    std = clean.std(ddof=1)
    sharpe = float((clean.mean() / std) * np.sqrt(12)) if std > 0 else np.nan
    return {
        "return": total_return,
        "cagr": cagr,
        "max_drawdown": float(drawdown.min()),
        "sharpe": sharpe,
    }


def smh_monthly_returns() -> pd.DataFrame:
    prices = pd.read_csv(PRICE_CACHE, index_col=0, header=[0, 1], parse_dates=True).sort_index()
    open_prices = prices["Open"]
    close_prices = prices["Close"]
    trading_days = close_prices.index
    rows = []
    periods = pd.Series(trading_days).dt.to_period("M").unique()
    for month_index, period in enumerate(periods):
        days = trading_days[trading_days.to_period("M") == period]
        trade_date = days[0]
        prior = trading_days[trading_days < trade_date]
        if not len(prior) or trade_date.date() < date(2010, 1, 1) or trade_date.date() > AS_OF:
            continue
        next_period = periods[month_index + 1] if month_index < len(periods) - 1 else None
        if next_period is None:
            exit_price = close_prices.loc[close_prices.index >= trade_date, "SMH"].dropna().iloc[-1]
        else:
            next_days = trading_days[trading_days.to_period("M") == next_period]
            exit_price = open_prices.loc[next_days[0], "SMH"]
        monthly_return = float(exit_price) / float(open_prices.loc[trade_date, "SMH"]) - 1.0
        rows.append(
            {
                "month": str(period),
                "trade_date": trade_date,
                "period": "IS" if trade_date.date() < date(2020, 1, 1) else "OOS",
                "monthly_return": monthly_return,
            }
        )
    return pd.DataFrame(rows)


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(text, style=style)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_indent(table, 120)
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        set_cell_shading(hdr_cells[idx], header_fill)
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in hdr_cells[idx].paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=9.5, color=INK, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=9.25, color=INK)
    set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "SMH Momentum Strategy Reference"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Prepared by Codex - backtest uses public SEC holdings and adjusted Yahoo Finance prices"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=8.5, color=MUTED)
    return doc


def add_masthead(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(TITLE)
    set_run_font(run, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(SUBTITLE)
    set_run_font(run, size=13.5, color=MUTED)

    metadata = [
        ("Prepared for", "Future strategy replication and review"),
        ("Data window", "2010-01-01 through 2026-05-24 YTD"),
        ("Selected strategy", "Top2 L252 S0 SMA100 DCA1"),
        ("Benchmark", "SMH buy-and-hold, monthly open-to-open"),
        ("Status", "Research backtest, not investment advice"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, color=INK, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    paragraph_border_bottom(rule, color="2E74B5", size="10")


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_indent(table, 120)
    set_table_borders(table, BORDER)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=INK)
    set_table_widths(table, [6.5])
    doc.add_paragraph()


def build() -> None:
    monthly = pd.read_excel(RESULTS_XLSX, sheet_name="Top3 monthly details")
    monthly = monthly[monthly["strategy"] == STRATEGY].copy()
    smh = smh_monthly_returns()
    monthly["year"] = pd.to_datetime(monthly["trade_date"]).dt.year
    smh["year"] = pd.to_datetime(smh["trade_date"]).dt.year

    periods = {
        "In-sample 2010-2019": monthly["period"] == "IS",
        "Out-of-sample 2020-2026 YTD": monthly["period"] == "OOS",
        "Full period 2010-2026 YTD": pd.Series(True, index=monthly.index),
    }
    performance_rows = []
    risk_rows = []
    for label, mask in periods.items():
        strategy_metrics = metric_series(monthly.loc[mask, "monthly_return"])
        smh_mask = smh["period"].eq("IS") if "In-sample" in label else smh["period"].eq("OOS") if "Out-of-sample" in label else pd.Series(True, index=smh.index)
        smh_metrics = metric_series(smh.loc[smh_mask, "monthly_return"])
        performance_rows.append(
            [
                label,
                format_percent(strategy_metrics["return"]),
                format_percent(smh_metrics["return"]),
                format_percent(strategy_metrics["return"] - smh_metrics["return"]),
                format_percent(strategy_metrics["cagr"]),
                format_percent(smh_metrics["cagr"]),
            ]
        )
        risk_rows.append(
            [
                label,
                format_percent(strategy_metrics["max_drawdown"]),
                format_percent(smh_metrics["max_drawdown"]),
                format_float(strategy_metrics["sharpe"]),
                format_float(smh_metrics["sharpe"]),
            ]
        )

    annual_rows = []
    for year in sorted(set(monthly["year"]) & set(smh["year"])):
        strategy_return = float((1.0 + monthly.loc[monthly["year"] == year, "monthly_return"]).prod() - 1.0)
        smh_return = float((1.0 + smh.loc[smh["year"] == year, "monthly_return"]).prod() - 1.0)
        annual_rows.append([str(year), format_percent(strategy_return), format_percent(smh_return), format_percent(strategy_return - smh_return)])

    cash_months = int((monthly["tickers"] == "CASH").sum())
    invested_months = int((monthly["tickers"] != "CASH").sum())
    ticker_counter: Counter[str] = Counter()
    for tickers in monthly.loc[monthly["tickers"] != "CASH", "tickers"]:
        for ticker in str(tickers).split(","):
            ticker_counter[ticker.strip()] += 1
    top_ticker_text = ", ".join(f"{ticker} ({count})" for ticker, count in ticker_counter.most_common(8))

    coverage = pd.read_excel(RESULTS_XLSX, sheet_name="Price coverage")
    unavailable = coverage[(coverage["ticker"] != "SMH") & (~coverage["available"])]["ticker"].tolist()

    doc = setup_document()
    add_masthead(doc)

    add_callout(
        doc,
        "One-page takeaway",
        "The selected strategy is a monthly, two-stock momentum rotation inside the historical SMH/HOLDRS universe. "
        "It beat SMH in the 2010-2019 in-sample window, the full 2010-2026 YTD test, and the 2020-2026 YTD OOS window on total return and drawdown. "
        "OOS Sharpe was slightly higher than SMH in this run.",
    )

    doc.add_heading("Selected Configuration", level=1)
    add_table(
        doc,
        ["Field", "Setting"],
        [
            ["Universe", "Historical SMH and legacy Semiconductor HOLDRS holdings from public SEC filings."],
            ["Point-in-time rule", "Use only the latest holdings snapshot filed on or before the monthly signal date."],
            ["Ranking score", "252-trading-day total momentum: Close[t] / Close[t-252] - 1."],
            ["Skip", "S0: no one-month skip. The score uses the signal date close."],
            ["Selection", "Buy the top 2 ranked stocks, equal weight, when risk filter is on."],
            ["Risk filter", "Risk-on only when SMH close is above its 100-day SMA at the signal date. Otherwise hold cash."],
            ["Execution", "Signal at prior month-end close; enter at next month first open; exit/rebalance at following month first open."],
            ["DCA1", "No staged entry. Move to full target exposure immediately on each monthly rebalance."],
            ["Benchmark", "SMH buy-and-hold, measured monthly open-to-open on the same calendar."],
        ],
        [1.55, 4.95],
    )

    doc.add_heading("Trading Rules", level=1)
    for item in [
        "At each month-end signal date, load the latest SMH/HOLDRS holdings filing that was public by that date.",
        "Drop stocks without usable adjusted open/close price history for that month.",
        "Calculate each eligible stock's 252-trading-day momentum using adjusted closes.",
        "Check the market filter: SMH adjusted close must be above its 100-day simple moving average.",
        "If the filter is on, buy the top 2 momentum stocks at the next trading day's adjusted open, equal weighted.",
        "If the filter is off, hold cash for the month.",
        "At the next monthly rebalance, repeat the process and replace the portfolio if the top 2 names change.",
    ]:
        add_numbered(doc, item)

    doc.add_heading("Performance Summary", level=1)
    add_table(
        doc,
        ["Period", "Strategy Return", "SMH Return", "Excess", "Strategy CAGR", "SMH CAGR"],
        performance_rows,
        [1.65, 1.0, 0.95, 0.85, 1.0, 0.95],
    )
    add_table(
        doc,
        ["Period", "Strategy Max DD", "SMH Max DD", "Strategy Sharpe", "SMH Sharpe"],
        risk_rows,
        [1.8, 1.1, 1.1, 1.2, 1.1],
    )

    doc.add_heading("Annual Return Comparison", level=1)
    add_paragraph(
        doc,
        "Annual rows are calendar-year compounded monthly returns. The 2026 row is year-to-date through 2026-05-24.",
    )
    add_table(
        doc,
        ["Year", "Strategy", "SMH", "Excess"],
        annual_rows,
        [0.8, 1.4, 1.4, 1.4],
    )

    doc.add_heading("Operational Notes", level=1)
    for item in [
        f"Months tested: {len(monthly)}. Invested months: {invested_months}. Cash months: {cash_months}.",
        f"Most frequently selected tickers in the selected strategy: {top_ticker_text}.",
        "Backtest uses adjusted Yahoo Finance open/close data. Adjusted prices include split/dividend adjustment behavior from the data source.",
        "Sharpe ratio is calculated from monthly returns, annualized by sqrt(12), with a zero risk-free-rate assumption.",
        "The test does not include commissions, bid/ask spread, market impact, taxes, or position-size constraints.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Data and Validation Notes", level=1)
    for item in [
        "Historical universe source: public SEC filings for legacy Semiconductor HOLDRS, Market Vectors Semiconductor ETF, and VanEck Semiconductor ETF.",
        "Lookahead control: holdings are keyed by filing date, not only by period end. The selected ticker audit found no cases where a trade used a filing after the signal date.",
        "Before December 2011, SMH was represented by the legacy Semiconductor HOLDRS structure. After the VanEck exchange-offer transition, the universe switches to VanEck/Market Vectors SMH filings.",
        f"Unavailable acquired/delisted tickers in this Yahoo-based run: {', '.join(unavailable)}.",
        "Because several old constituents lack usable Yahoo price data, this is more robust than a current-holdings test but still not a fully institutional survivorship-free database.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Source Files", level=1)
    source_rows = [
        ["Backtest workbook", str(RESULTS_XLSX.relative_to(ROOT))],
        ["Historical holdings file", "data/smh_components/smh_historical_holdings_sec.csv"],
        ["Backtest script", "src/strategy_lab/smh_historical_components_momentum_is_oos.py"],
        ["SEC holdings extractor", "src/strategy_lab/extract_smh_historical_holdings_sec.py"],
    ]
    add_table(doc, ["Artifact", "Path"], source_rows, [1.7, 4.8])

    doc.core_properties.title = TITLE
    doc.core_properties.subject = SUBTITLE
    doc.core_properties.keywords = "SMH, momentum, backtest, point-in-time holdings"
    doc.core_properties.author = "Codex"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
